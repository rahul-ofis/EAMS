from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime
import io

def create_word_document(goals_data, employee_names):
    doc = Document()
    
    # Create title
    title = doc.add_heading('Team Goals Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    for employee_id, goals in goals_data.items():
        if not goals:
            continue
            
        # Add employee name as heading
        doc.add_heading(f'Employee: {employee_names.get(employee_id, "Unknown")}', 1)
        
        for goal in goals:
            # Goal description
            p = doc.add_paragraph()
            p.add_run('Goal Description: ').bold = True
            p.add_run(goal['description'])
            
            # Goal KPIs
            p = doc.add_paragraph()
            p.add_run('KPIs: ').bold = True
            p.add_run(goal['kpis'])
            
            # Weightage
            p = doc.add_paragraph()
            p.add_run('Weightage: ').bold = True
            p.add_run(f"{goal['weightage']}%")
            
            # Manager Rating
            if goal.get('manager_rating'):
                p = doc.add_paragraph()
                p.add_run('Manager Rating: ').bold = True
                rating_run = p.add_run(str(goal['manager_rating']))
                rating_run.font.color.rgb = RGBColor(0, 112, 192)
            
            # Manager Feedback
            if goal.get('manager_feedback'):
                p = doc.add_paragraph()
                p.add_run('Manager Feedback: ').bold = True
                feedback_run = p.add_run(goal['manager_feedback'])
                feedback_run.font.italic = True
            
            doc.add_paragraph('---')  # Separator
            
        doc.add_page_break()
    
    # Save to memory
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def create_excel_sheet(goals_data, employee_names):
    # Create a Pandas Excel writer
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook = writer.book
        
        # Create formats
        header_format = workbook.add_format({
            'bold': True,
            'font_size': 12,
            'bg_color': '#4472C4',
            'font_color': 'white',
            'border': 1
        })
        
        cell_format = workbook.add_format({
            'font_size': 11,
            'border': 1
        })
        
        # Create sheets for each employee
        for employee_id, goals in goals_data.items():
            if not goals:
                continue
                
            employee_name = employee_names.get(employee_id, "Unknown")
            df = pd.DataFrame(goals)
            
            # Reorder and rename columns
            columns = ['description', 'kpis', 'weightage', 'manager_rating', 'manager_feedback']
            column_names = ['Goal Description', 'KPIs', 'Weightage (%)', 'Manager Rating', 'Manager Feedback']
            df = df[columns]
            df.columns = column_names
            
            # Write to sheet
            sheet_name = employee_name[:31]  # Excel sheet names limited to 31 chars
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            # Format the sheet
            worksheet = writer.sheets[sheet_name]
            
            # Set column widths
            worksheet.set_column('A:A', 40)  # Goal Description
            worksheet.set_column('B:B', 40)  # KPIs
            worksheet.set_column('C:C', 15)  # Weightage
            worksheet.set_column('D:D', 15)  # Rating
            worksheet.set_column('E:E', 40)  # Feedback
            
            # Apply formats
            for col_num, _ in enumerate(df.columns):
                worksheet.write(0, col_num, df.columns[col_num], header_format)
            
            # Apply cell format to data
            for row_num in range(len(df)):
                for col_num in range(len(df.columns)):
                    worksheet.write(row_num + 1, col_num, df.iloc[row_num, col_num], cell_format)
    
    output.seek(0)
    return output
